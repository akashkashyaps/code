import os
import re
import docx
from tqdm import tqdm
from langchain_ollama import ChatOllama
from langchain.schema import HumanMessage

class AgenticReportGrader:
    def __init__(self,
                 base_directory: str,
                 model: str = 'mistral:7b-instruct',
                 num_ctx: int = 2048,
                 temperature: float = 0.7,
                 top_p: float = 0.9,
                 num_predict: int = 3000):
        self.base_directory = base_directory
        self.llm = ChatOllama(
            model=model,
            num_ctx=num_ctx,
            temperature=temperature,
            top_p=top_p,
            num_predict=num_predict
        )

    def llm_call(self, prompt: str) -> str:
        """
        Call the language model with a direct prompt, returning the model's text output.
        """
        messages = [HumanMessage(content=prompt)]
        response = self.llm(messages)
        return response.content if hasattr(response, "content") else response

    def _extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract all text from a DOCX file, concatenating paragraphs with newlines.
        """
        doc = docx.Document(file_path)
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())

    def _parse_rubric_text(self, text: str):
        """
        Parse the rubric string to identify:
         1. 'Grading Criteria' text (if present).
         2. Each numbered section of the form:
            '1. Title (XX%)' with everything after it up to the next heading or EOS.
         Returns:
          {
            "grading_criteria": "text under 'Grading Criteria' heading if found",
            "sections": [
               {
                 "section_number": int,
                 "section_heading": str,   # e.g. "1. Introduction... (10%)"
                 "title": str,            # e.g. "Introduction... (10%)"
                 "weight": float,
                 "body_text": str
               }, ...
            ]
          }
        """

        # 1) Extract the “Grading Criteria” block if it exists.
        #    We'll look for a line that has "Grading Criteria" and capture everything
        #    until the first numbered heading or the end of text.
        grading_criteria_pattern = re.compile(
            r'(Grading Criteria.*?)(?=\n\s*\d+\.\s.*?\(\d+%\)|$)',
            flags=re.IGNORECASE|re.DOTALL
        )
        match_gc = grading_criteria_pattern.search(text)
        grading_criteria_text = match_gc.group(1).strip() if match_gc else ""

        # 2) Extract each numbered section heading of the form:
        #    e.g. "1. Introduction: (10%)"   capturing:
        #      - the entire heading
        #      - the number
        #      - the weight
        #
        #    We also want to grab everything until the next heading or end of text.
        #    Explanation of pattern:
        #    - Look for <digit+>.  (like "1.")
        #    - Some text, eventually containing (XX%)
        #    - Then capture all text (body) until the next heading or end
        section_pattern = re.compile(
            r'(?P<header>(?P<section_number>\d+)\.\s*(?P<title>.*?\(\s*(?P<weight>\d+)\s*%\)))(?P<body>.*?)(?=\n\s*\d+\.\s.*?\(\d+%\)|$)',
            flags=re.DOTALL
        )

        sections = []
        for m in section_pattern.finditer(text):
            section_num = int(m.group('section_number'))
            header = m.group('header').strip()
            title_str = m.group('title').strip()
            weight_str = m.group('weight').strip()
            weight_val = float(weight_str)
            body = m.group('body').strip()

            sections.append({
                "section_number": section_num,
                "section_heading": header,
                "title": title_str,
                "weight": weight_val,
                "body_text": body
            })

        return {
            "grading_criteria": grading_criteria_text,
            "sections": sections
        }

    def _build_section_prompts(self, parsed_rubric) -> list:
        """
        Build an individual section prompt for each parsed rubric section.
        Display them as well so you can see the “section prompts.”
        """
        prompts = []
        for section in parsed_rubric["sections"]:
            # Construct a concise or detailed prompt:
            section_prompt = f"""
Section {section['section_number']} – {section['section_heading']}

Full Rubric Subsection Text:
{section['body_text']}

Instructions:
1. Evaluate the student's work specifically on this section's criteria.
2. Provide a numeric or letter grade for this section.
3. Provide rationale in 1-2 short paragraphs.
"""
            section_prompt = section_prompt.strip()

            prompts.append(section_prompt)
            # Display prompt
            print(f"\n--- SECTION PROMPT {section['section_number']} ---")
            print(section_prompt)
            print("---------------------------------------------")
        return prompts

    def _build_final_prompt(self, parsed_rubric, section_evaluations: list) -> str:
        """
        Create a single final prompt that references:
         - The entire 'Grading Criteria' text (if needed).
         - Each section’s weight.
         - Summaries or references to each section's evaluation result.
         - Instructions to compute final score from these weights.
        Display it so you can see the “final prompt.”
        """
        # Summarize the sections with their weights
        sections_summary = []
        for s, ev in zip(parsed_rubric["sections"], section_evaluations):
            sections_summary.append(
                f"Section {s['section_number']} ({s['weight']}%): {s['title']}\n"
                f"Evaluation Summary:\n{ev}\n"
            )
        sections_text = "\n".join(sections_summary)

        final_prompt = f"""
You have evaluated each section individually. Now combine those results into a final score.

Grading Criteria Reference:
{parsed_rubric["grading_criteria"]}

Section Evaluations (with assigned weights):
{sections_text}

Instructions:
1. For each section, note the numeric grade you assigned (out of 100) or letter grade.
2. Convert each section’s grade into a numeric score, multiply by the weight (%).
3. Sum these weighted scores to get the final overall score.
4. Provide the final result, along with a rationale referencing how well the work met all criteria.
"""
        final_prompt = final_prompt.strip()

        # Display final prompt
        print("\n=== FINAL PROMPT ===")
        print(final_prompt)
        print("====================\n")

        return final_prompt

    def evaluate_section(self, section_prompt: str, report_text: str, section_number: int) -> str:
        """
        Evaluate a single section using the model, returning the section-specific evaluation.
        """
        prompt = f"""
{section_prompt}

--- Student Report Below ---
{report_text}
"""
        return self.llm_call(prompt.strip())

    def grade_reports(self):
        for folder_name in tqdm(os.listdir(self.base_directory), desc="Processing"):
            folder_path = os.path.join(self.base_directory, folder_name)
            if not os.path.isdir(folder_path):
                continue

            # Identify a student report file in the folder:
            report_files = [
                f for f in os.listdir(folder_path)
                if f.lower().startswith('report_') and f.endswith('.docx')
            ]
            if not report_files:
                continue
            report_path = os.path.join(folder_path, report_files[0])
            report_text = self._extract_text_from_docx(report_path)

            # Identify the prompt_#.docx files:
            for prompt_num in range(1, 7):
                prompt_files = [
                    f for f in os.listdir(folder_path)
                    if f.lower() == f'prompt_{prompt_num}.docx'
                ]
                if not prompt_files:
                    continue
                prompt_path = os.path.join(folder_path, prompt_files[0])
                rubric_text = self._extract_text_from_docx(prompt_path)

                try:
                    # 1. Parse the rubric to get grading criteria and sections.
                    parsed_rubric = self._parse_rubric_text(rubric_text)

                    # 2. Build & display prompts for each section.
                    section_prompts = self._build_section_prompts(parsed_rubric)

                    # 3. Evaluate each section using the student's report.
                    section_evaluations = []
                    for sp, s_info in zip(section_prompts, parsed_rubric["sections"]):
                        eval_text = self.evaluate_section(sp, report_text, s_info["section_number"])
                        section_evaluations.append(eval_text)

                    # 4. Build the final prompt referencing all partial evaluations + weighting
                    final_prompt = self._build_final_prompt(parsed_rubric, section_evaluations)

                    # 5. Call the model with the final prompt to produce a final combined result
                    final_output = self.llm_call(final_prompt)

                    # 6. Save final output to a DOCX file
                    output_filename = f'GRADED_{report_files[0].replace(".docx", "")}_Prompt_{prompt_num}.docx'
                    output_path = os.path.join(folder_path, output_filename)
                    doc = docx.Document()
                    doc.add_paragraph(final_output)
                    doc.save(output_path)

                except Exception as e:
                    print(f"Error processing {folder_name}/Prompt_{prompt_num}: {e}")

def main():
    base_directory = '/home/akash/Downloads/grading_documents'
    models = {
        'qwen2.5:7b-instruct-q4_0': 32768,
        'deepseek-r1:7b-qwen-distill-q4_K_M': 32768,
        'deepseek-r1:8b-llama-distill-q4_K_M': 32768
    }

    for model, ctx in models.items():
        print(f"\nStarting grading with {model}")
        grader = AgenticReportGrader(
            base_directory=base_directory,
            model=model,
            num_ctx=ctx,
            temperature=0.1,
            top_p=0.9,
            num_predict=3000
        )
        grader.grade_reports()
        print(f"Completed grading with {model}")

if __name__ == '__main__':
    main()
