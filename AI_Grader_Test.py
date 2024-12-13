import os
import docx
from tqdm import tqdm
from langchain.llms import Ollama

class ReportGrader:
    def __init__(
        self,
        base_directory: str,
        model: str = 'mistral:7b-instruct',
        num_ctx: int = 2048,
        temperature: float = 0.7,
        top_p: float = 0.9
    ):
        self.base_directory = base_directory
        self.model = model
        self.num_ctx = num_ctx
        self.temperature = temperature
        self.top_p = top_p

        # Initialize the Ollama LLM with specified settings
        self.llm = Ollama(
            model=self.model,
            num_ctx=self.num_ctx,
            temperature=self.temperature,
            top_p=self.top_p,
            # Add other model parameters as needed
        )

    def _extract_text_from_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        
        raw_text = '\n'.join([para.text for para in doc.paragraphs if para.text])
        
        processed_text = raw_text.lower()
        for term in ['gpt', 'openai', 'chatgpt']:
            processed_text = processed_text.replace(term, '')
        
        return processed_text

    def grade_reports(self):
        for folder_name in tqdm(os.listdir(self.base_directory), desc=f"Processing Folders with {self.model}"):
            folder_path = os.path.join(self.base_directory, folder_name)
            
            if not os.path.isdir(folder_path):
                continue
            
            # Find the report file
            report_files = [f for f in os.listdir(folder_path) 
                            if f.lower().startswith('report_') and f.lower().endswith('.docx')]
            
            if not report_files:
                continue
            
            # Process the report
            report_path = os.path.join(folder_path, report_files[0])
            report_text = self._extract_text_from_docx(report_path)
            report_name = os.path.splitext(report_files[0])[0]
            
            # Process each prompt from 1 to 6
            for prompt_num in range(1, 7):
                # More flexible prompt file matching
                prompt_files = [f for f in os.listdir(folder_path) 
                                if f.lower() == f'prompt_{prompt_num}.docx']
                
                if not prompt_files:
                    continue
                
                prompt_path = os.path.join(folder_path, prompt_files[0])
                prompt_text = self._extract_text_from_docx(prompt_path)
                
                try:
                    # Use the prompt text as-is, combining it with the report content
                    complete_prompt = f"{prompt_text}\n\n{report_text}"

                    # Generate response from Ollama
                    response = self.llm(complete_prompt)

                    # Prepare output filename
                    output_filename = f'{self.model}_{report_name}_Prompt_{prompt_num}.docx'
                    output_path = os.path.join(folder_path, output_filename)
                    
                    # Save response to a new Word document
                    output_doc = docx.Document()
                    
                    # Add the original prompt to the document for reference
                    output_doc.add_paragraph(f"Original Prompt (Prompt_{prompt_num}):")
                    output_doc.add_paragraph(prompt_text)
                    output_doc.add_paragraph("\n--- AI Response ---\n")
                    
                    output_doc.add_paragraph(response)
                    output_doc.save(output_path)
                    
                    print(f"Processed {folder_name} - Prompt {prompt_num} with {self.model}")
                
                except Exception as e:
                    print(f"Error processing {folder_name}/Prompt_{prompt_num} with {self.model}: {e}")

def main():
    base_directory = '/home/akash/Downloads/grading_documents'

    # Dictionary of models with their specific context lengths
    models = {
        'qwen2.5:7b-instruct-q4_0': 32768, 
        'llama3.1:8b-instruct-q4_0': 131072, 
        'gemma2:9b-instruct-q4_0': 8192, 
        'internlm2:7b': 32768,
        'mistral-nemo:12b-instruct-2407-q4_0': 1024000
    }

    # Run with each model
    for model, context_length in models.items():
        print(f"\nStarting processing with model: {model}")
        
        # Create grader with current model and its specific context length
        grader = ReportGrader(
            base_directory, 
            model=model, 
            num_ctx=context_length,
            temperature=0.5,
            top_p=0.9
        )
        
        # Run grading process for this model
        grader.grade_reports()
        
        print(f"Completed processing with model: {model}")

if __name__ == '__main__':
    main()